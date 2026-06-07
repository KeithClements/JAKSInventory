"""Convert PHASE_1_TEST_PLAN.md (or any markdown) to .docx using python-docx."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_formatted_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    # Inline: **bold**, *italic*, `code`, [text](url) -> text
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        chunk = m.group(0)
        if chunk.startswith("**"):
            run = p.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*"):
            run = p.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`"):
            run = p.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif chunk.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if link:
                p.add_run(link.group(1))
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells if c)


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    i = 0
    in_code = False
    code_lines: list[str] = []
    list_buffer: list[tuple[str, str]] = []  # (kind, text) kind = ul|ol|bq

    def flush_list():
        nonlocal list_buffer
        for kind, text in list_buffer:
            if kind == "bq":
                _add_formatted_paragraph(doc, text, style="Intense Quote")
            elif kind == "ul":
                _add_formatted_paragraph(doc, text, style="List Bullet")
            elif kind == "ol":
                _add_formatted_paragraph(doc, text, style="List Number")
        list_buffer = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                flush_list()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            flush_list()
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.LINE)
            i += 1
            continue

        # Table
        if line.strip().startswith("|") and i + 1 < len(lines) and "|" in lines[i + 1]:
            flush_list()
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = _parse_table_row(lines[i])
                if not _is_separator_row(cells):
                    table_rows.append(cells)
                i += 1
            if table_rows:
                tbl = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                tbl.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci, cell_text in enumerate(row):
                        if ci < len(tbl.rows[ri].cells):
                            cell = tbl.rows[ri].cells[ci]
                            cell.text = re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text)
                            if ri == 0:
                                for p in cell.paragraphs:
                                    for r in p.runs:
                                        r.bold = True
                                _set_cell_shading(cell, "E8E8E8")
                doc.add_paragraph()
            continue

        # Headings
        hm = re.match(r"^(#{1,4})\s+(.+)$", line)
        if hm:
            flush_list()
            level = len(hm.group(1))
            text = hm.group(2)
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            styles = {1: "Title", 2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}
            style = styles.get(level, "Heading 3")
            if level == 1 and doc.paragraphs:
                style = "Heading 1"
            p = doc.add_paragraph(text, style=style)
            if level == 1 and not doc.paragraphs[0].text:
                pass
            i += 1
            continue

        # Blockquote
        if line.strip().startswith(">"):
            list_buffer.append(("bq", line.strip().lstrip(">").strip()))
            i += 1
            continue

        # Unordered list
        if re.match(r"^(\s*)[-*]\s+", line):
            m = re.match(r"^\s*[-*]\s+(.*)$", line)
            if m:
                list_buffer.append(("ul", m.group(1)))
            i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            m = re.match(r"^\s*\d+\.\s+(.*)$", line)
            if m:
                list_buffer.append(("ol", m.group(1)))
            i += 1
            continue

        # Checkbox lines
        if re.match(r"^-\s+\[[ xX]?\]", line):
            checked = "[x]" in line.lower() or "[X]" in line
            text = re.sub(r"^-\s+\[[ xX]?\]\s*", "", line)
            sym = "☑ " if checked else "☐ "
            list_buffer.append(("ul", sym + text))
            i += 1
            continue

        if line.strip() == "":
            flush_list()
            i += 1
            continue

        flush_list()
        _add_formatted_paragraph(doc, line.strip())
        i += 1

    flush_list()
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else root / "PHASE_1_TEST_PLAN.md"
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix(".docx")
    convert_md_to_docx(md, out)
